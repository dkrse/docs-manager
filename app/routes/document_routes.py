import os
import io
import html as html_module
import zipfile
import hashlib
import hmac
import secrets
import tempfile
from datetime import datetime, date
from fastapi import APIRouter, Depends, Request, Form, UploadFile, File, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, JSONResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from sqlalchemy import or_, func
from app.database import get_db
from app.models import User, Document, UserSettings, Favorite
from app.auth import get_current_user
from app.config import UPLOAD_DIR, SECRET_KEY, MAX_UPLOAD_SIZE
import unicodedata
import fitz  # PyMuPDF
import markdown as md
from docx import Document as DocxDocument
import openpyxl
import csv


def _doc_sign(doc_id: int) -> str:
    today = date.today().isoformat()
    return hmac.new(SECRET_KEY.encode(), f"{doc_id}:{today}".encode(), hashlib.sha256).hexdigest()[:12]


def make_doc_hash(doc_id: int) -> str:
    """Generate a daily-rotating hash: hex(id)-signature."""
    return f"{doc_id:x}-{_doc_sign(doc_id)}"


def resolve_doc_hash(token: str, db: Session) -> Document | None:
    """Resolve hash token to document. O(1) lookup."""
    parts = token.rsplit("-", 1)
    if len(parts) != 2:
        return None
    try:
        doc_id = int(parts[0], 16)
    except ValueError:
        return None
    if _doc_sign(doc_id) != parts[1]:
        return None
    return db.query(Document).filter(Document.id == doc_id).first()

from datetime import timezone as _tz

_LOCAL_TZ = datetime.now().astimezone().tzinfo


def _local_dt(dt_val, fmt="%Y-%m-%d %H:%M") -> str:
    if not dt_val:
        return ""
    if dt_val.tzinfo is None:
        dt_val = dt_val.replace(tzinfo=_tz.utc)
    return dt_val.astimezone(_LOCAL_TZ).strftime(fmt)


router = APIRouter()
templates = Jinja2Templates(directory="app/templates")

THUMBNAILS_DIR = UPLOAD_DIR / "thumbnails"
THUMBNAILS_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx", ".xlsx", ".csv"}

FILE_TYPE_ICONS = {
    ".pdf": "bi-file-earmark-pdf text-danger",
    ".md": "bi-file-earmark-text text-info",
    ".txt": "bi-file-earmark-text text-secondary",
    ".docx": "bi-file-earmark-word text-primary",
    ".xlsx": "bi-file-earmark-excel text-success",
    ".csv": "bi-file-earmark-spreadsheet text-success",
}


def extract_text(file_path: str, extension: str) -> str:
    """Extract plain text content from a document file."""
    ext = extension.lower()
    try:
        if ext == ".pdf":
            doc = fitz.open(file_path)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        elif ext == ".docx":
            doc = DocxDocument(file_path)
            parts = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        elif ext in (".md", ".txt"):
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        elif ext == ".xlsx":
            with open(file_path, "rb") as fxl:
                wb = openpyxl.load_workbook(fxl, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts.append(" ".join(str(c) for c in row if c is not None))
            wb.close()
            return "\n".join(parts)
        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                return "\n".join(" ".join(row) for row in reader)
    except Exception:
        return ""
    return ""


def _safe_path(file_path: str) -> bool:
    """Verify file_path is within UPLOAD_DIR."""
    try:
        return os.path.realpath(file_path).startswith(str(UPLOAD_DIR))
    except (ValueError, OSError):
        return False


def _escape_like(s: str) -> str:
    """Escape SQL LIKE wildcards."""
    return s.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")


def remove_diacritics(s: str) -> str:
    nfkd = unicodedata.normalize('NFKD', s)
    return ''.join(c for c in nfkd if not unicodedata.combining(c))


def generate_thumbnail(pdf_path: str, page_num: int, thumb_path: str):
    doc = fitz.open(pdf_path)
    page_idx = max(0, min(page_num - 1, len(doc) - 1))
    page = doc[page_idx]
    mat = fitz.Matrix(2, 2)
    pix = page.get_pixmap(matrix=mat)
    pix.save(thumb_path)
    doc.close()


def get_user_settings(user: User, db: Session) -> UserSettings:
    settings = db.query(UserSettings).filter(UserSettings.user_id == user.id).first()
    if not settings:
        settings = UserSettings(user_id=user.id)
        db.add(settings)
        db.commit()
        db.refresh(settings)
    return settings


import re

def _esc(text: str) -> str:
    return html_module.escape(text)


async def _save_upload(file: UploadFile) -> tuple[str, str, int]:
    """Stream upload to temp file, compute hash. Returns (temp_path, sha256_hex, file_size). Raises ValueError if too large."""
    h = hashlib.sha256()
    size = 0
    tmp = tempfile.NamedTemporaryFile(delete=False, dir=str(UPLOAD_DIR))
    try:
        while True:
            chunk = await file.read(64 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > MAX_UPLOAD_SIZE:
                tmp.close()
                os.unlink(tmp.name)
                raise ValueError(f"File exceeds maximum size of {MAX_UPLOAD_SIZE // (1024*1024)}MB")
            h.update(chunk)
            tmp.write(chunk)
        tmp.close()
    except ValueError:
        raise
    except Exception:
        tmp.close()
        os.unlink(tmp.name)
        raise
    return tmp.name, h.hexdigest(), size


def _sanitize_html(html_str: str) -> str:
    """Remove script/iframe/object/embed tags from HTML output."""
    html_str = re.sub(r'<\s*script[^>]*>.*?</\s*script\s*>', '', html_str, flags=re.DOTALL | re.IGNORECASE)
    html_str = re.sub(r'<\s*script[^>]*/?\s*>', '', html_str, flags=re.IGNORECASE)
    html_str = re.sub(r'<\s*/?\s*(iframe|object|embed|form)[^>]*>', '', html_str, flags=re.IGNORECASE)
    html_str = re.sub(r'\bon\w+\s*=', '', html_str, flags=re.IGNORECASE)
    return html_str


def xlsx_to_html(file_path: str) -> str:
    with open(file_path, "rb") as fxl:
        wb = openpyxl.load_workbook(fxl, data_only=True)
    html_parts = []
    for ws in wb.worksheets:
        html_parts.append(f"<h5>{_esc(ws.title)}</h5>")
        html_parts.append("<table class='table table-bordered table-sm table-striped'>")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            tag = "th" if i == 0 else "td"
            html_parts.append("<tr>")
            for cell in row:
                html_parts.append(f"<{tag}>{_esc(str(cell)) if cell is not None else ''}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")
    wb.close()
    return "\n".join(html_parts)


def csv_to_html(file_path: str) -> str:
    html_parts = ["<table class='table table-bordered table-sm table-striped'>"]
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            tag = "th" if i == 0 else "td"
            html_parts.append("<tr>")
            for cell in row:
                html_parts.append(f"<{tag}>{_esc(cell)}</{tag}>")
            html_parts.append("</tr>")
    html_parts.append("</table>")
    return "\n".join(html_parts)


def docx_to_html(file_path: str) -> str:
    doc = DocxDocument(file_path)
    html_parts = []
    for para in doc.paragraphs:
        style = para.style.name.lower() if para.style else ""
        text = para.text
        if not text.strip():
            html_parts.append("<br>")
            continue
        if "heading 1" in style:
            html_parts.append(f"<h1>{_esc(text)}</h1>")
        elif "heading 2" in style:
            html_parts.append(f"<h2>{_esc(text)}</h2>")
        elif "heading 3" in style:
            html_parts.append(f"<h3>{_esc(text)}</h3>")
        elif "heading" in style:
            html_parts.append(f"<h4>{_esc(text)}</h4>")
        else:
            runs_html = ""
            for run in para.runs:
                t = _esc(run.text) if run.text else ""
                if not t:
                    continue
                if run.bold:
                    t = f"<strong>{t}</strong>"
                if run.italic:
                    t = f"<em>{t}</em>"
                if run.underline:
                    t = f"<u>{t}</u>"
                runs_html += t
            if runs_html:
                html_parts.append(f"<p>{runs_html}</p>")
            else:
                html_parts.append(f"<p>{_esc(text)}</p>")

    for table in doc.tables:
        html_parts.append("<table class='table table-bordered table-sm'>")
        for i, row in enumerate(table.rows):
            html_parts.append("<tr>")
            tag = "th" if i == 0 else "td"
            for cell in row.cells:
                html_parts.append(f"<{tag}>{_esc(cell.text)}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    return "\n".join(html_parts)


@router.get("/api/documents", response_class=JSONResponse)
async def api_documents(
    request: Request,
    search: str = "",
    category: str = "",
    file_type: str = "",
    sort: str = "",
    page: int = 1,
    doc_date_from: str = "",
    doc_date_to: str = "",
    upload_date_from: str = "",
    upload_date_to: str = "",
    favorites_only: str = "",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    settings = get_user_settings(current_user, db)
    page_size = settings.page_size or 20
    if not sort:
        sort = settings.default_sort or "random"

    query = db.query(Document)

    if not current_user.is_admin:
        query = query.filter(
            or_(
                Document.uploaded_by == current_user.id,
                Document.is_private == False,
            )
        )

    hidden = settings.hidden_hashtags.strip() if settings.hidden_hashtags else ""
    if hidden:
        hidden_tags = [t.strip().lower() for t in hidden.split(",") if t.strip()]
        for tag in hidden_tags:
            query = query.filter(~func.lower(Document.hashtags).ilike(f"%{_escape_like(tag)}%", escape="\\"))

    if favorites_only == "1":
        fav_doc_ids = [r[0] for r in db.query(Favorite.document_id).filter(Favorite.user_id == current_user.id).all()]
        query = query.filter(Document.id.in_(fav_doc_ids))

    if search:
        normalized = remove_diacritics(search.lower())
        like = f"%{_escape_like(normalized)}%"
        query = query.filter(
            or_(
                func.lower(Document.original_filename).ilike(like, escape="\\"),
                func.lower(Document.description).ilike(like, escape="\\"),
                func.lower(Document.notes).ilike(like, escape="\\"),
                func.lower(Document.hashtags).ilike(like, escape="\\"),
                func.lower(Document.content).ilike(like, escape="\\"),
            )
        )
    if category:
        query = query.filter(Document.category == category)
    if file_type:
        query = query.filter(Document.file_extension == file_type)

    if doc_date_from:
        try:
            query = query.filter(Document.document_date >= datetime.strptime(doc_date_from, "%Y-%m-%d"))
        except ValueError:
            pass
    if doc_date_to:
        try:
            query = query.filter(Document.document_date <= datetime.strptime(doc_date_to, "%Y-%m-%d").replace(hour=23, minute=59, second=59))
        except ValueError:
            pass
    if upload_date_from:
        try:
            query = query.filter(Document.uploaded_at >= datetime.strptime(upload_date_from, "%Y-%m-%d"))
        except ValueError:
            pass
    if upload_date_to:
        try:
            query = query.filter(Document.uploaded_at <= datetime.strptime(upload_date_to, "%Y-%m-%d").replace(hour=23, minute=59, second=59))
        except ValueError:
            pass

    sort_map = {
        "upload_desc": Document.uploaded_at.desc(),
        "upload_asc": Document.uploaded_at.asc(),
        "date_desc": Document.document_date.desc(),
        "date_asc": Document.document_date.asc(),
        "size_desc": Document.file_size.desc(),
        "size_asc": Document.file_size.asc(),
        "name_asc": func.lower(Document.original_filename).asc(),
        "name_desc": func.lower(Document.original_filename).desc(),
        "type_asc": Document.file_extension.asc(),
        "type_desc": Document.file_extension.desc(),
        "modified_desc": Document.updated_at.desc(),
        "modified_asc": Document.updated_at.asc(),
        "random": func.random(),
    }
    if sort == "random":
        query = query.order_by(func.random())
    else:
        order = sort_map.get(sort, Document.uploaded_at.desc())
        query = query.order_by(order, Document.id)

    total = query.count()
    total_pages = max(1, (total + page_size - 1) // page_size)
    page = max(1, min(page, total_pages))
    documents = query.offset((page - 1) * page_size).limit(page_size).all()

    uploader_ids = list(set(doc.uploaded_by for doc in documents))
    user_cache = {}
    if uploader_ids:
        users = db.query(User).filter(User.id.in_(uploader_ids)).all()
        user_cache = {u.id: u.username for u in users}
    fav_ids = set(
        r[0] for r in db.query(Favorite.document_id).filter(Favorite.user_id == current_user.id).all()
    )
    results = []
    for doc in documents:
        results.append({
            "id": doc.id,
            "original_filename": doc.original_filename,
            "file_extension": doc.file_extension or ".pdf",
            "description": doc.description or "",
            "category": doc.category or "",
            "hashtags": doc.hashtags or "",
            "is_private": doc.is_private,
            "uploaded_by": doc.uploaded_by,
            "uploader_name": user_cache.get(doc.uploaded_by, "Unknown"),
            "uploaded_at": _local_dt(doc.uploaded_at),
            "updated_at": _local_dt(doc.updated_at),
            "has_thumbnail": bool(doc.thumbnail_path),
            "notes": doc.notes or "",
            "file_size": doc.file_size or 0,
            "document_date": doc.document_date.strftime("%Y-%m-%d") if doc.document_date else "",
            "is_favorite": doc.id in fav_ids,
            "doc_hash": make_doc_hash(doc.id),
        })

    return {
        "documents": results,
        "page": page,
        "page_size": page_size,
        "total_pages": total_pages,
        "total": total,
        "is_admin": current_user.is_admin,
        "user_id": current_user.id,
        "sort": sort,
        "show_edit": settings.show_edit if settings.show_edit is not None else True,
        "show_download": settings.show_download if settings.show_download is not None else True,
        "show_delete": settings.show_delete if settings.show_delete is not None else True,
        "show_line_numbers": settings.show_line_numbers if settings.show_line_numbers is not None else False,
    }


@router.get("/", response_class=HTMLResponse)
async def dashboard(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    all_categories = [r[0] for r in db.query(Document.category).distinct().all() if r[0]]
    settings = get_user_settings(current_user, db)
    return templates.TemplateResponse("dashboard.html", {
        "request": request,
        "user": current_user,
        "categories": all_categories,
        "default_sort": settings.default_sort or "random",
    })


@router.get("/upload", response_class=HTMLResponse)
async def upload_page(request: Request, current_user: User = Depends(get_current_user)):
    if current_user.is_admin:
        raise HTTPException(status_code=403, detail="Admins cannot upload documents")
    return templates.TemplateResponse("upload.html", {"request": request, "user": current_user, "error": None})


@router.post("/upload", response_class=HTMLResponse)
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    category: str = Form(""),
    hashtags: str = Form(""),
    description: str = Form(""),
    notes: str = Form(""),
    is_private: bool = Form(False),
    document_date: str = Form(""),
    thumbnail_page: int = Form(1),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.is_admin:
        raise HTTPException(status_code=403, detail="Admins cannot upload documents")

    ext = os.path.splitext(file.filename)[1].lower() if file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return templates.TemplateResponse("upload.html", {
            "request": request, "user": current_user,
            "error": f"Allowed file types: {', '.join(ALLOWED_EXTENSIONS)}"
        })

    try:
        tmp_path, file_hash, file_size = await _save_upload(file)
    except ValueError as e:
        return templates.TemplateResponse("upload.html", {
            "request": request, "user": current_user, "error": str(e)
        })

    existing = db.query(Document).filter(Document.content_hash == file_hash).first()
    if existing:
        os.unlink(tmp_path)
        return templates.TemplateResponse("upload.html", {
            "request": request, "user": current_user,
            "error": f"This file already exists as '{existing.original_filename}'"
        })

    file_path = UPLOAD_DIR / file_hash
    os.rename(tmp_path, str(file_path))

    thumb_path = None
    if ext == ".pdf":
        thumb_filename = f"{file_hash}.png"
        thumb_path = THUMBNAILS_DIR / thumb_filename
        try:
            generate_thumbnail(str(file_path), thumbnail_page, str(thumb_path))
        except Exception:
            thumb_path = None

    doc_date = None
    if document_date:
        try:
            doc_date = datetime.strptime(document_date, "%Y-%m-%d")
        except ValueError:
            pass

    extracted_text = remove_diacritics(extract_text(str(file_path), ext).lower())

    doc = Document(
        filename=file_hash,
        original_filename=file.filename,
        file_extension=ext,
        file_path=str(file_path),
        file_size=file_size,
        content_hash=file_hash,
        thumbnail_path=str(thumb_path) if thumb_path else None,
        thumbnail_page=thumbnail_page if ext == ".pdf" else 1,
        category=category.strip(),
        hashtags=hashtags.strip(),
        description=description.strip(),
        notes=notes.strip(),
        is_private=is_private,
        uploaded_by=current_user.id,
        document_date=doc_date,
        content=extracted_text,
    )
    db.add(doc)
    db.commit()
    return RedirectResponse(url="/", status_code=302)


@router.get("/document/{doc_id}/thumbnail")
async def get_thumbnail(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404)
    if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403)
    if doc.thumbnail_path and os.path.exists(doc.thumbnail_path):
        return FileResponse(doc.thumbnail_path, media_type="image/png")
    raise HTTPException(status_code=404)


@router.get("/document/{doc_id}/info", response_class=JSONResponse)
async def document_info(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404)
    if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403)
    uploader = db.query(User).filter(User.id == doc.uploaded_by).first()
    return {
        "id": doc.id,
        "original_filename": doc.original_filename,
        "file_extension": doc.file_extension or ".pdf",
        "description": doc.description or "",
        "category": doc.category or "",
        "hashtags": doc.hashtags or "",
        "notes": doc.notes or "",
        "is_private": doc.is_private,
        "uploader_name": uploader.username if uploader else "Unknown",
        "uploaded_at": _local_dt(doc.uploaded_at),
        "document_date": doc.document_date.strftime("%Y-%m-%d") if doc.document_date else "",
        "file_size": doc.file_size or 0,
        "share_token": doc.share_token or "",
        "doc_hash": make_doc_hash(doc.id),
    }


@router.get("/document/{doc_id}/download")
async def download_document(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    if not _safe_path(doc.file_path):
        raise HTTPException(status_code=403, detail="Access denied")
    return FileResponse(doc.file_path, filename=doc.original_filename, media_type="application/octet-stream")


@router.get("/document/{doc_id}/pdf")
async def serve_pdf(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    return FileResponse(doc.file_path, media_type="application/pdf")


@router.get("/document/{doc_id}/render", response_class=HTMLResponse)
async def render_document(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """Render markdown, txt, or docx as HTML for the modal viewer."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404)
    if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403)

    ext = (doc.file_extension or "").lower()
    body = ""

    if ext == ".md":
        with open(doc.file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        body = _sanitize_html(md.markdown(raw, extensions=["tables", "fenced_code", "codehilite"]))
    elif ext == ".txt":
        with open(doc.file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        body = f"<pre style='white-space:pre-wrap; word-wrap:break-word;'>{_esc(raw)}</pre>"
    elif ext == ".docx":
        try:
            body = docx_to_html(doc.file_path)
        except Exception as e:
            body = f"<p class='text-danger'>Error rendering document: {_esc(str(e))}</p>"
    elif ext == ".xlsx":
        try:
            body = xlsx_to_html(doc.file_path)
        except Exception as e:
            body = f"<p class='text-danger'>Error rendering spreadsheet: {_esc(str(e))}</p>"
    elif ext == ".csv":
        try:
            body = csv_to_html(doc.file_path)
        except Exception as e:
            body = f"<p class='text-danger'>Error rendering CSV: {_esc(str(e))}</p>"
    else:
        raise HTTPException(status_code=400, detail="Unsupported format for rendering")

    html_page = f"""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<link href="/static/css/bootstrap.min.css" rel="stylesheet">
<style>body {{ padding: 20px; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
pre {{ background: #f8f9fa; padding: 15px; border-radius: 4px; }}
table {{ margin: 10px 0; }}
img {{ max-width: 100%; }}</style>
</head><body>{body}</body></html>"""
    return HTMLResponse(html_page)


@router.get("/v/{doc_hash}/pdf")
async def view_pdf_by_hash(doc_hash: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = resolve_doc_hash(doc_hash, db)
    if not doc:
        raise HTTPException(status_code=404)
    if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403)
    return FileResponse(doc.file_path, media_type="application/pdf")


@router.get("/v/{doc_hash}/render", response_class=HTMLResponse)
async def view_render_by_hash(doc_hash: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = resolve_doc_hash(doc_hash, db)
    if not doc:
        raise HTTPException(status_code=404)
    if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403)
    ext = (doc.file_extension or "").lower()
    body = ""
    if ext == ".md":
        with open(doc.file_path, "r", encoding="utf-8", errors="replace") as f:
            body = _sanitize_html(md.markdown(f.read(), extensions=["tables", "fenced_code", "codehilite"]))
    elif ext == ".txt":
        with open(doc.file_path, "r", encoding="utf-8", errors="replace") as f:
            body = f"<pre style='white-space:pre-wrap; word-wrap:break-word;'>{_esc(f.read())}</pre>"
    elif ext == ".docx":
        try:
            body = docx_to_html(doc.file_path)
        except Exception as e:
            body = f"<p class='text-danger'>Error: {_esc(str(e))}</p>"
    elif ext == ".xlsx":
        try:
            body = xlsx_to_html(doc.file_path)
        except Exception as e:
            body = f"<p class='text-danger'>Error: {_esc(str(e))}</p>"
    elif ext == ".csv":
        try:
            body = csv_to_html(doc.file_path)
        except Exception as e:
            body = f"<p class='text-danger'>Error: {_esc(str(e))}</p>"
    else:
        raise HTTPException(status_code=400)
    html_page = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<link href="/static/css/bootstrap.min.css" rel="stylesheet">
<style>body {{ padding: 20px; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
pre {{ background: #f8f9fa; padding: 15px; border-radius: 4px; }}
table {{ margin: 10px 0; }} img {{ max-width: 100%; }}</style>
</head><body>{body}</body></html>"""
    return HTMLResponse(html_page)


@router.post("/document/{doc_id}/delete")
async def delete_document(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_admin and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    if doc.thumbnail_path and os.path.exists(doc.thumbnail_path):
        os.remove(doc.thumbnail_path)
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/", status_code=302)


@router.get("/api/search-context", response_class=JSONResponse)
async def search_context(
    search: str = "",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not search or len(search) < 1:
        return {"results": []}

    settings = get_user_settings(current_user, db)
    normalized = remove_diacritics(search.lower())
    like = f"%{_escape_like(normalized)}%"

    query = db.query(Document)
    if not current_user.is_admin:
        query = query.filter(
            or_(Document.uploaded_by == current_user.id, Document.is_private == False)
        )

    query = query.filter(
        or_(
            func.lower(Document.original_filename).ilike(like, escape="\\"),
            func.lower(Document.description).ilike(like, escape="\\"),
            func.lower(Document.notes).ilike(like, escape="\\"),
            func.lower(Document.hashtags).ilike(like, escape="\\"),
            func.lower(Document.content).ilike(like, escape="\\"),
        )
    )

    docs = query.limit(50).all()
    results = []
    for doc in docs:
        # Find first matching line in content
        match_line = ""
        match_field = ""
        for field_name, field_val in [
            ("content", doc.content or ""),
            ("filename", doc.original_filename or ""),
            ("description", doc.description or ""),
            ("notes", doc.notes or ""),
            ("hashtags", doc.hashtags or ""),
        ]:
            if not field_val:
                continue
            norm_val = remove_diacritics(field_val.lower()) if field_name == "content" else field_val.lower()
            if normalized in norm_val:
                # Find the actual line
                lines = field_val.split("\n") if field_name == "content" else [field_val]
                for line in lines:
                    norm_line = remove_diacritics(line.lower()) if field_name == "content" else line.lower()
                    if normalized in norm_line:
                        match_line = line.strip()
                        match_field = field_name
                        break
                if match_line:
                    break

        results.append({
            "id": doc.id,
            "original_filename": doc.original_filename,
            "file_extension": doc.file_extension or "",
            "match_line": match_line[:500],
            "match_field": match_field,
            "doc_hash": make_doc_hash(doc.id),
        })

    return {"results": results, "query": search}


@router.post("/api/bulk-delete", response_class=JSONResponse)
async def bulk_delete(request: Request, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    data = await request.json()
    ids = data.get("ids", [])[:200]
    docs = db.query(Document).filter(Document.id.in_(ids)).all() if ids else []
    deleted = 0
    for doc in docs:
        if not current_user.is_admin and doc.uploaded_by != current_user.id:
            continue
        if _safe_path(doc.file_path) and os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        if doc.thumbnail_path and _safe_path(doc.thumbnail_path) and os.path.exists(doc.thumbnail_path):
            os.remove(doc.thumbnail_path)
        db.delete(doc)
        deleted += 1
    db.commit()
    return {"ok": True, "deleted": deleted}


@router.post("/api/bulk-zip")
async def bulk_zip(request: Request, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    data = await request.json()
    ids = data.get("ids", [])[:200]  # limit to 200 documents
    docs = db.query(Document).filter(Document.id.in_(ids)).all() if ids else []
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    tmp_path = tmp.name
    tmp.close()
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for doc in docs:
            if not current_user.is_admin and doc.is_private and doc.uploaded_by != current_user.id:
                continue
            if _safe_path(doc.file_path) and os.path.exists(doc.file_path):
                zf.write(doc.file_path, doc.original_filename)
    from starlette.background import BackgroundTask
    return FileResponse(tmp_path, media_type="application/zip", filename="documents.zip",
                        background=BackgroundTask(os.unlink, tmp_path))


@router.post("/api/document/{doc_id}/share", response_class=JSONResponse)
async def toggle_share(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404)
    if not current_user.is_admin and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403)
    if doc.share_token:
        doc.share_token = None
        db.commit()
        return {"ok": True, "share_token": ""}
    token = secrets.token_urlsafe(32)
    doc.share_token = token
    db.commit()
    return {"ok": True, "share_token": token}


@router.get("/shared/{token}")
async def public_shared_view(token: str, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.share_token == token).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Link not found or expired")
    return FileResponse(doc.file_path, filename=doc.original_filename, media_type="application/octet-stream")


@router.get("/shared/{token}/view", response_class=HTMLResponse)
async def public_shared_viewer(token: str, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.share_token == token).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Link not found or expired")
    ext = (doc.file_extension or "").lower()
    safe_name = _esc(doc.original_filename)
    if ext == ".pdf":
        return HTMLResponse(f"""<!DOCTYPE html><html><head><meta charset="utf-8"><title>{safe_name}</title>
        <link href="/static/css/bootstrap.min.css" rel="stylesheet">
        </head><body class="p-3">
        <h5>{safe_name}</h5>
        <iframe src="/shared/{_esc(token)}" width="100%" style="height:85vh;border:none;"></iframe>
        <a href="/shared/{_esc(token)}" class="btn btn-sm btn-success mt-2"><i class="bi bi-download"></i> Download</a>
        </body></html>""")
    body = ""
    if ext == ".md":
        with open(doc.file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        body = md.markdown(raw, extensions=["tables", "fenced_code"])
        body = _sanitize_html(body)
    elif ext == ".txt":
        with open(doc.file_path, "r", encoding="utf-8", errors="replace") as f:
            body = f"<pre style='white-space:pre-wrap;'>{_esc(f.read())}</pre>"
    elif ext == ".docx":
        body = docx_to_html(doc.file_path)
    elif ext == ".xlsx":
        body = xlsx_to_html(doc.file_path)
    elif ext == ".csv":
        body = csv_to_html(doc.file_path)
    return HTMLResponse(f"""<!DOCTYPE html><html><head><meta charset="utf-8"><title>{safe_name}</title>
    <link href="/static/css/bootstrap.min.css" rel="stylesheet">
    <style>body {{ padding: 20px; }}</style>
    </head><body><h5>{safe_name}</h5>{body}
    <hr><a href="/shared/{_esc(token)}" class="btn btn-sm btn-success"><i class="bi bi-download"></i> Download</a>
    </body></html>""")


@router.post("/api/document/{doc_id}/favorite", response_class=JSONResponse)
async def toggle_favorite(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404)
    existing = db.query(Favorite).filter(Favorite.user_id == current_user.id, Favorite.document_id == doc_id).first()
    if existing:
        db.delete(existing)
        db.commit()
        return {"ok": True, "is_favorite": False}
    db.add(Favorite(user_id=current_user.id, document_id=doc_id))
    db.commit()
    return {"ok": True, "is_favorite": True}


@router.post("/api/upload", response_class=JSONResponse)
async def api_upload_document(
    request: Request,
    file: UploadFile = File(...),
    category: str = Form(""),
    hashtags: str = Form(""),
    description: str = Form(""),
    notes: str = Form(""),
    is_private: bool = Form(False),
    document_date: str = Form(""),
    thumbnail_page: int = Form(1),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.is_admin:
        return JSONResponse({"ok": False, "error": "Admins cannot upload documents"}, status_code=403)

    ext = os.path.splitext(file.filename)[1].lower() if file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return JSONResponse({"ok": False, "error": f"Unsupported file type: {ext}", "filename": file.filename})

    try:
        tmp_path, file_hash, file_size = await _save_upload(file)
    except ValueError as e:
        return JSONResponse({"ok": False, "error": str(e), "filename": file.filename})

    existing = db.query(Document).filter(Document.content_hash == file_hash).first()
    if existing:
        os.unlink(tmp_path)
        return JSONResponse({"ok": False, "error": f"Already exists as '{existing.original_filename}'", "filename": file.filename})

    file_path = UPLOAD_DIR / file_hash
    os.rename(tmp_path, str(file_path))

    thumb_path = None
    if ext == ".pdf":
        thumb_filename = f"{file_hash}.png"
        thumb_path = THUMBNAILS_DIR / thumb_filename
        try:
            generate_thumbnail(str(file_path), thumbnail_page, str(thumb_path))
        except Exception:
            thumb_path = None

    doc_date = None
    if document_date:
        try:
            doc_date = datetime.strptime(document_date, "%Y-%m-%d")
        except ValueError:
            pass

    extracted_text = remove_diacritics(extract_text(str(file_path), ext).lower())

    doc = Document(
        filename=file_hash,
        original_filename=file.filename,
        file_extension=ext,
        file_path=str(file_path),
        file_size=file_size,
        content_hash=file_hash,
        thumbnail_path=str(thumb_path) if thumb_path else None,
        thumbnail_page=thumbnail_page if ext == ".pdf" else 1,
        category=category.strip(),
        hashtags=hashtags.strip(),
        description=description.strip(),
        notes=notes.strip(),
        is_private=is_private,
        uploaded_by=current_user.id,
        document_date=doc_date,
        content=extracted_text,
    )
    db.add(doc)
    db.commit()
    return JSONResponse({"ok": True, "filename": file.filename})


# --- Edit metadata route ---

@router.post("/document/{doc_id}/edit", response_class=JSONResponse)
async def edit_document_metadata(
    request: Request,
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_admin and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    form = await request.form()
    doc.category = form.get("category", doc.category or "").strip()
    doc.hashtags = form.get("hashtags", doc.hashtags or "").strip()
    doc.description = form.get("description", doc.description or "").strip()
    doc.notes = form.get("notes", doc.notes or "").strip()
    is_private = form.get("is_private", "")
    doc.is_private = is_private == "true"
    doc_date_str = form.get("document_date", "").strip()
    if doc_date_str:
        try:
            doc.document_date = datetime.strptime(doc_date_str, "%Y-%m-%d")
        except ValueError:
            pass
    else:
        doc.document_date = None
    db.commit()
    return {"ok": True}


# --- Settings routes ---

@router.get("/help", response_class=HTMLResponse)
async def help_page(request: Request, current_user: User = Depends(get_current_user)):
    return templates.TemplateResponse("help.html", {"request": request, "user": current_user})


@router.get("/settings", response_class=HTMLResponse)
async def settings_page(request: Request, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    settings = get_user_settings(current_user, db)
    return templates.TemplateResponse("settings.html", {
        "request": request, "user": current_user, "settings": settings, "success": None
    })


@router.post("/settings", response_class=HTMLResponse)
async def save_settings(
    request: Request,
    page_size: int = Form(20),
    default_sort: str = Form("random"),
    hidden_hashtags: str = Form(""),
    theme: str = Form("light"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    settings = get_user_settings(current_user, db)
    settings.page_size = max(5, min(100, page_size))
    settings.default_sort = default_sort
    settings.hidden_hashtags = hidden_hashtags.strip()
    settings.theme = theme if theme in ("light", "dark") else "light"
    db.commit()
    response = templates.TemplateResponse("settings.html", {
        "request": request, "user": current_user, "settings": settings, "success": "Settings saved"
    })
    response.set_cookie("theme", settings.theme, httponly=False, samesite="lax", max_age=31536000)
    return response


@router.get("/api/settings", response_class=JSONResponse)
async def api_get_settings(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    s = get_user_settings(current_user, db)
    return {
        "theme": s.theme or "light",
        "page_size": s.page_size or 20,
        "default_sort": s.default_sort or "random",
        "hidden_hashtags": s.hidden_hashtags or "",
        "show_edit": s.show_edit if s.show_edit is not None else True,
        "show_download": s.show_download if s.show_download is not None else True,
        "show_delete": s.show_delete if s.show_delete is not None else True,
        "show_line_numbers": s.show_line_numbers if s.show_line_numbers is not None else False,
    }


@router.post("/api/settings", response_class=JSONResponse)
async def api_save_settings(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    form = await request.form()
    s = get_user_settings(current_user, db)
    s.theme = form.get("theme", "light")
    if s.theme not in ("light", "dark"):
        s.theme = "light"
    try:
        s.page_size = max(5, min(100, int(form.get("page_size", 20))))
    except (ValueError, TypeError):
        s.page_size = 20
    s.default_sort = form.get("default_sort", "random")
    s.hidden_hashtags = form.get("hidden_hashtags", "").strip()
    s.show_edit = form.get("show_edit", "true") == "true"
    s.show_download = form.get("show_download", "true") == "true"
    s.show_delete = form.get("show_delete", "true") == "true"
    s.show_line_numbers = form.get("show_line_numbers", "false") == "true"
    db.commit()
    resp = JSONResponse({"ok": True})
    resp.set_cookie("theme", s.theme, httponly=False, samesite="lax", max_age=31536000)
    return resp
